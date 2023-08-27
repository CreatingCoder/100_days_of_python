from docx import Document





my_file = open('names.txt', 'r')       
info = my_file.read()
file_list = info.split("\n")
my_file.close()
        

for name in file_list:

    doc = Document()

    doc.add_paragraph(f'{name},')


    #add paragraph
    doc.add_paragraph('I am writing this letter to express my interest in blah blah blah')

    #add paragraph
    doc.add_paragraph('I am writing this letter to express my interest in blah blah blah')

    #add paragraph
    doc.add_paragraph('I am writing this letter to express my interest in blah blah blah')

    #add paragraph
    doc.add_paragraph('V/R,')
    doc.add_paragraph('User User')
    
    
    #save document
    doc.save('C:\\Users\\Shane\\OneDrive\\Documents\\Python\\100DOP\\day_23\\docx_app\\Ready_to_send\\' + f'Letter for {name}.docx')


