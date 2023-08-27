from docx import Document
from datetime import date


username = 'John Doe'
address_1 = '123 Place Drive'
address_2 = 'Baltimore, CA 12345'
address_3 = ''
to_block = 'Recipient'
position = 'Computer Science'

phone = '(240) 750-5232'
email = 'asdfasdf@adsf.com'
years_exp = 4
job_num =  '123456asdfadsf'


# username = input('Please enter your name ')
# address_1  = input('Please enter your address line 1')
# address_2 = input('Please enter your address line 2')
# address_3 = input('Please enter your address line 3 (if applicable) ')
# to_block = input('Please enter the title/name the letter is being sent to: ')
# position = input('Please enter the job position: ')


today = date.today()
today_formatted = today.strftime("%d/%m/%Y")

doc = Document()

#\n {address_3}
doc.add_paragraph(f'{username}\n{address_1}\n{address_2}\n{phone}')

doc.add_paragraph(f'{today_formatted}')

doc.add_paragraph(f'Dear {to_block}')

doc.add_paragraph(f'I am writing to express my interest in the position of {position} with the job announcement number of {job_num}.  With over {years_exp} years experience in the {position} field, I believe that I would be a good fit for this position.  Futhermore, I have ')

doc.add_paragraph(f'For your immediate review, I have attached my current resume.  If you have any questions or comments, I can be reached via telephone at {phone} or email at {email}')

doc.add_paragraph(f'V/R,')

doc.add_paragraph(f'{username}')


##I saw the posting on







#Save last to save all the changes made
doc.save('test.docx')
